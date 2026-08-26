"""
Loaders: PDF, DOCX, image -> Blocks. Guide module M4.

DECISION FOR THIS SESSION, from your 80/20 argument: DETECT, DO NOT STITCH.

Detection needs all four signals (column count, column x-positions, a table
touching the bottom margin on page N, a failed header check on page N+1) and is
most of the work. Joining two row-lists is a handful of lines. So flagging is not
deferring the effort -- it is building the part needed either way and stopping
short of the part that can manufacture data.

The asymmetry that settles it:
  MISS a continuation  -> an incomplete answer from text that really exists.
                          0.712 instead of 0.831. Bad, and inspectable.
  INVENT a continuation -> a table present in NO source document: page 1's
                          headers above page 2's unrelated numbers, clean
                          structure, uniform columns, a real page citation.
                          Every structural validator passes it, because its
                          structure is perfect.

Same shape as the zip() truncation and the cache-dimension bug: the dangerous
failure is the one that produces well-formed output.

---------------------------------------------------------------------------
WHAT THIS FILE MEASURES, AND WHY TWO NUMBERS

`n_continuation_suspects` alone is circular -- it counts what the detector
finds. strategy="lines" cannot see a whitespace-aligned table at all (the
fixture's Table 5 becomes prose), so a low count may be a blind spot rather than
a rare phenomenon. The detector's count is therefore reported next to a
hand-labelled sample (`n_tables_undetected_manual`), the same way ANN recall is
only meaningful against exact search.

---------------------------------------------------------------------------
TABLE HANDLING IN SESSION 1

Strategy (a): tables are serialised to markdown inline by pymupdf4llm, and that
is the whole table pipeline for now. find_tables() runs SEPARATELY as a detector
and validator only -- it emits diagnostics and flags, never extra blocks. That
avoids indexing every table twice (duplicates crowd top-k, which is Barnett
failure point 7) and keeps strategies (b) and (c) for Session 3.

Two measured facts drive that split:
  - find_tables(strategy="lines") is high precision, low recall: no ruling
    lines means no table.
  - find_tables(strategy="text") is the opposite, and catastrophically so: it
    returned page 1's title and prose as a 23x3 table with words cut mid-token.
  - find_tables().extract() MANGLES underscores -- 'HYBRID_RRF' comes back as
    'HYBRID RRF' plus a stray '_' -- while page.get_text() on the same page
    returns it verbatim. Cell text is lower fidelity than page text, and exact
    identifiers are precisely what BM25 exists to match (M2). So cell text is
    never the indexed text here; it is only ever evidence for the validator.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

# MODULE-LEVEL AND MANDATORY, and this line is the whole fix for a measurement
# bug that moved aggregate numbers by 5x.
#
# pymupdf4llm calls pymupdf.TOOLS.unset_quad_corrections(True) at MODULE level,
# so merely IMPORTING it changes glyph quad geometry -- and therefore changes
# find_tables() results. Measured on lost-in-the-middle.pdf: 3 tables before the
# import, 9 after. Other files move the other way (hnsw 11 -> 2, graphrag 9 -> 5)
# because it is a geometry change: some tables appear, others vanish.
#
# It is also STICKY: calling unset_quad_corrections(False) afterwards does not
# restore the original count. So this cannot be controlled with a flag, only with
# import order.
#
# Which means the loader's regime is the ONLY regime, and any script that
# measures the corpus without this import measures a DIFFERENT PARSER. That is
# exactly what happened -- a standalone detection sweep imported only pymupdf and
# reported 61 tables / 11 headerless, while the real loader reports different
# numbers for the same files. Neither was wrong; they were two parsers.
#
# So: import it here, unconditionally, before anything can call find_tables, and
# name it in PARSER_VERSION so provenance records which parser produced an index.
# An undeclared global input that changes output while content_hash and
# parser_version stay identical is precisely what parsed_hash exists to catch.
import pymupdf  # noqa: F401
import pymupdf4llm  # noqa: F401  -- imported for its global side effect, not its API

from .. import config, limits
from .document import (
    Block as _Block,
    Block,
    ChunkKind,
    DocType,
    Source,
    sha256_text,
)

PARSER_VERSION = "pymupdf4llm@1.28+lines-detector@1+quadcorr-unset+cleanhdr@2+pagetextrepair@3+bboxflags@1"

# --------------------------------------------------------------------------
# Parse cache -- the real budget constraint on this corpus
# --------------------------------------------------------------------------
#
# Measured, not guessed: table detection alone is 102 seconds over 137 pages
# (hnsw.pdf is 35s for 13 pages), and the full pymupdf4llm markdown pass over
# the corpus blew a seven-minute timeout. Quota was the assumed scarce resource;
# PARSING WALL-CLOCK is worse. Uncached, every chunking experiment for the rest
# of the project pays that again -- which is how a sweep across five configs
# stops being runnable.
#
# THE KEY IS THE INPUTS, NOT THE OUTPUT. content_hash + PARSER_VERSION. It is
# tempting to key on parsed_hash, and that cannot work: parsed_hash is computed
# FROM the parse, so you would have to parse in order to look up the parse. The
# two fields have different jobs --
#     key  (content_hash + parser_version) -> "have I done this work?"
#     value fingerprint (parsed_hash)       -> "did the parser drift?"
# and the second is only meaningful because it is not the first.
PARSE_CACHE = config.CACHE_DIR / "parse"


def _parse_cache_path(content_hash: str) -> Path:
    h = sha256_text(f"{PARSER_VERSION}|{content_hash}")
    return PARSE_CACHE / h[:2] / f"{h}.json"


def _cache_load(src: Source) -> tuple[Source, list[Block], dict[str, Any]] | None:
    p = _parse_cache_path(src.content_hash)
    if not p.exists():
        return None
    try:
        d = json.loads(p.read_text(encoding="utf-8"))
        blocks = [Block.from_json(b) for b in d["blocks"]]
        src.parsed_hash = d["parsed_hash"]
        src.n_pages = d.get("n_pages", src.n_pages)
        diag = d["diagnostics"]
        diag["from_cache"] = True
        return src, blocks, diag
    except Exception:  # noqa: BLE001 -- a corrupt entry is a miss, never a crash
        return None


def _cache_store(src: Source, blocks: list[Block], diag: dict[str, Any]) -> None:
    p = _parse_cache_path(src.content_hash)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".tmp")
    tmp.write_text(
        json.dumps(
            {
                "parser_version": PARSER_VERSION,
                "content_hash": src.content_hash,
                "parsed_hash": src.parsed_hash,
                "n_pages": src.n_pages,
                "blocks": [b.to_json() for b in blocks],
                "diagnostics": {k: v for k, v in diag.items() if k != "from_cache"},
            },
            indent=1,
        ),
        encoding="utf-8",
    )
    os.replace(tmp, p)

# Bottom-margin proximity below which a table is a continuation candidate.
# 80pt is roughly one inch: close enough that the table plausibly ran out of
# page. Tuned against the fixture (its page-1 table sits at 70pt) and reported
# as a knob rather than buried, because it is exactly the sort of threshold that
# quietly decides a measurement.
BOTTOM_MARGIN_PT = 80.0

# Column x-positions must agree within this to count as "same columns".
COL_X_TOLERANCE_PT = 6.0


# --------------------------------------------------------------------------
# Table validation -- the check, borrowed from strategy (c) without its machinery
# --------------------------------------------------------------------------


def _squash(s: str) -> str:
    """Alphanumerics only, lowercased.

    Whitespace-stripping alone is not enough, and finding that out cost a round:
    find_tables() reports the cell as "HYBRID RRF _" while page.get_text()
    renders "HYBRID_RRF". Squashing whitespace gives "hybridrrf_" vs "hybrid_rrf"
    -- still unequal, because the two extractors disagree about where PUNCTUATION
    goes, not just whitespace. Comparing across extractors has to be insensitive
    to both.
    """
    return re.sub(r"[^a-z0-9]", "", str(s or "").lower())


def normalize_cell(c: Any) -> str:
    """Collapse ALL whitespace inside a cell, newlines included.

    Not cosmetic. find_tables().extract() splits characters onto their own lines:
    "HYBRID_RRF" arrives as "HYBRID RRF" + newline + "_", and -- far worse --
    "0.831" arrives as "0831" + newline + ".". Without collapsing, the header
    check asks float("0831
.") which raises, concludes the cell is not numeric,
    and passes the headerless fragment. That is how v2 of this check still failed
    on the case it was written for: correct rule, defeated by dirty input.

    Note what the same mangling does to the VALUE: "0831." parses as 831.0, not
    0.831 -- wrong by three orders of magnitude, in a table with flawless
    structure. This is the decisive argument for never indexing cell text.
    Strategy (c) (extract to a real table, query it) would have loaded every
    number in the fixture off by 1000x and passed validation.
    """
    return " ".join(str(c or "").split())


def first_nonempty(rows: list[list[Any]]) -> list[str]:
    """The bug fix that mattered.

    Version one of the header check read rows[0], which on the fixture's page-2
    fragment is ['', '', '', ''] -- zero numeric cells, so it returned PASS on
    the exact case it existed to catch. Right about WHAT to inspect, wrong about
    WHERE.
    """
    for r in rows:
        cells = [normalize_cell(c) for c in r]
        if any(cells):
            return cells
    return []


def _is_numeric(s: str) -> bool:
    t = "".join(str(s).split()).replace(",", "").replace("%", "").rstrip(".")
    if not t:
        return False
    try:
        float(t)
        return True
    except ValueError:
        return False


def header_plausible(rows: list[list[Any]]) -> tuple[bool, str]:
    """Does the first real row look like LABELS rather than VALUES?

    Uniform cell count -- the obvious check -- cannot work here. The fixture's
    headerless fragment has three cells per row throughout, so 'a header exists
    and every row has the same width' passes it and the bug ships. The defect is
    semantic, not structural: a header's cells are labels, and labels are not
    numbers.

    Kept deliberately simple. It is a smoke alarm, not a fire marshal: it must
    catch the fixture and must not reject Table 4 (a legitimate table whose
    header is 'Parameter | Value | Applies at | Notes'). A cleverer check that
    rejects real tables is worse than this one, because an eval harness that
    excludes legitimate tables biases itself optimistically -- your point that a
    flag stops being risk-free once something acts on it.
    """
    hdr = first_nonempty(rows)
    if not hdr:
        return False, "no non-empty row"
    filled = [c for c in hdr if c]
    n_num = sum(_is_numeric(c) for c in filled)
    if n_num >= max(1, len(filled) // 2):
        return (
            False,
            f"{n_num}/{len(filled)} first-row cells are numeric -> data row, no header",
        )
    return True, f"header reads as labels: {hdr}"


def corrupt_numeric_cells(rows: list[list[Any]], page_text: str) -> list[tuple[str, str]]:
    """Cells whose text does NOT appear in the page's own text layer.

    THE FIX FOR THE 831.0 BUG, and note what it is not: it is not a smarter
    validator. It is a SECOND EXTRACTOR that disagrees.

    find_tables().extract() returned "0831." for a cell the document renders as
    "0.831" -- a thousand times too big, in a table with four columns, uniform
    rows, a real header and a real page number. Every structural check passes it.

    A range check cannot catch it: "recall <= 1.0" needs to know the column is
    recall, which needs the header, and 18% of detected tables have none. A
    distributional check cannot catch it either: the mangling is systematic, so
    the whole column shifts together -- [684.0, 601.0, 712.0, 831.0, 889.0] is
    internally consistent and entirely wrong.

    But page.get_text() renders "0.831" correctly on the same page. So the check
    is agreement between two extraction paths. Same principle as measuring ANN
    recall against exact search, and detector counts against hand labels: you
    cannot validate a measurement from inside it.
    """
    flat = " ".join(page_text.split())
    bad: list[tuple[str, str]] = []
    for row in rows:
        for cell in row:
            t = normalize_cell(cell)
            if not t or not _is_numeric(t):
                continue
            compact = t.replace(" ", "")
            if compact not in flat and t not in flat:
                bad.append((t, compact))
    return bad


def match_candidate(block_text: str, candidates: list["TableCandidate"]) -> "TableCandidate | None":
    """Which detected table is this markdown block?

    Content-based, never positional. pymupdf4llm's table traversal and
    find_tables()'s are two different orders, so pairing them by index is a
    guess, and a wrong pairing would attach one table's flags (or one table's
    repair text) to a different table.

    Used by BOTH the flag attribution and the text repair, deliberately: two
    matchers for the same question is how the flags and the repair would end up
    disagreeing about which table they were talking about.
    """
    blk = _squash(block_text)
    for cand in candidates:
        probes = [_squash(c) for c in cand.first_row if len(_squash(c)) >= 5]
        if probes and all(pr in blk for pr in probes[:2]):
            return cand
    return None


@dataclass
class TableCandidate:
    """One detected table plus everything the continuation check needs."""

    page: int
    n_rows: int
    n_cols: int
    bbox: tuple[float, float, float, float]
    col_xs: tuple[float, ...]
    bottom_gap: float
    top_gap: float
    header_ok: bool
    header_reason: str
    first_row: list[str] = field(default_factory=list)
    corrupt_cells: list[tuple[str, str]] = field(default_factory=list)
    # Set during the continuation pass, and carried on the CANDIDATE rather than
    # on a page number. Page-level attribution meant every block on a page with
    # one bad table inherited the flag -- so Table 4, with a perfectly good
    # header, reported table_header_missing=True. Harmless until something acts
    # on the field, and the eval is about to slice on it to ask "do repaired
    # tables answer worse?". Comparing a population that includes healthy tables
    # against one that does not is the text_source label bug one step earlier.
    is_continuation_suspect: bool = False
    clip: str = ""   # bbox-clipped page text, for the header-missing repair


def _column_xs(table: Any) -> tuple[float, ...]:
    """Column boundary x-positions, the strongest same-table signal available.

    Two fragments of one table share a column layout because they were laid out
    by the same code with the same widths. Two unrelated tables almost never do.
    """
    cells = getattr(table, "cells", None) or []
    xs = sorted({round(c[0], 1) for c in cells if c})
    if xs:
        return tuple(xs)
    x0, _, x1, _ = table.bbox
    return (round(x0, 1), round(x1, 1))


def detect_tables(page: Any, page_no: int) -> list[TableCandidate]:
    """strategy='lines' only, and that choice is a measured tradeoff.

    strategy='text' was tested and rejected: on the fixture it reported page 1's
    title and body prose as a 23x3 table, splitting words mid-token
    ('Meridian Ret' / 'rieval Benc'). High recall bought with precision so poor
    that every downstream validator becomes noise. The cost of 'lines' is a real
    blind spot -- an unruled table is simply not seen -- and that blind spot is
    why prevalence is reported against a hand-labelled sample rather than alone.
    """
    out: list[TableCandidate] = []
    try:
        found = page.find_tables(strategy="lines").tables
    except Exception:  # noqa: BLE001 -- a page that cannot be scanned is not a crash
        return out
    page_text = page.get_text()
    for t in found:
        rows = t.extract()
        ok, why = header_plausible(rows)
        corrupt = corrupt_numeric_cells(rows, page_text)
        out.append(
            TableCandidate(
                page=page_no,
                n_rows=t.row_count,
                n_cols=t.col_count,
                bbox=tuple(round(v, 1) for v in t.bbox),  # type: ignore[arg-type]
                col_xs=_column_xs(t),
                bottom_gap=round(page.rect.y1 - t.bbox[3], 1),
                top_gap=round(t.bbox[1] - page.rect.y0, 1),
                header_ok=ok,
                header_reason=why,
                first_row=first_nonempty(rows),
                corrupt_cells=corrupt,
            )
        )
    return out


def continuation_signals(
    prev: TableCandidate, cur: TableCandidate
) -> tuple[bool, list[str]]:
    """All four signals, reported individually rather than as one boolean.

    Returning the evidence matters: 'suspect' with no reasons is unauditable, and
    the whole justification for flagging over stitching is that a human decides.
    Requiring three of four keeps a single coincidence from raising a flag.
    """
    reasons: list[str] = []
    if cur.page == prev.page + 1:
        reasons.append("consecutive pages")
    if cur.n_cols == prev.n_cols:
        reasons.append(f"same column count ({cur.n_cols})")
    if len(cur.col_xs) == len(prev.col_xs) and all(
        abs(a - b) <= COL_X_TOLERANCE_PT for a, b in zip(cur.col_xs, prev.col_xs)
    ):
        reasons.append("same column x-positions")
    if prev.bottom_gap <= BOTTOM_MARGIN_PT:
        reasons.append(f"predecessor touches bottom margin ({prev.bottom_gap:.0f}pt)")
    if not cur.header_ok:
        reasons.append(f"no header on continuation: {cur.header_reason}")
    strong = ("consecutive pages" in reasons) and len(reasons) >= 3
    return strong, reasons


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def load_pdf(path: Path) -> tuple[Source, list[Block], dict[str, Any]]:
    """Prose+headings from pymupdf4llm; tables detected separately as diagnostics."""
    src = Source.from_file(path, DocType.PDF)

    doc = pymupdf.open(path)
    src.n_pages = doc.page_count

    # -- detection pass -------------------------------------------------------
    per_page: list[list[TableCandidate]] = [
        detect_tables(doc[i], i + 1) for i in range(doc.page_count)
    ]
    all_tables = [t for page in per_page for t in page]
    suspects: list[dict[str, Any]] = []
    for i in range(1, doc.page_count):
        for cur in per_page[i]:
            # A continuation starts at the top of its page; a table further down
            # had room above it and therefore did not run out of page.
            if cur.top_gap > 120:
                continue
            for prev in per_page[i - 1]:
                strong, reasons = continuation_signals(prev, cur)
                if strong:
                    cur.is_continuation_suspect = True
                    suspects.append(
                        {
                            "from_page": prev.page,
                            "to_page": cur.page,
                            "n_cols": cur.n_cols,
                            "first_row": cur.first_row,
                            "reasons": reasons,
                        }
                    )
    # REPAIR CANDIDATES, captured while the doc is still open.
    #
    # A header-missing table's markdown serialisation is not merely lossy, it is
    # WRONG: pymupdf4llm promotes the first data row to a column header, so
    # "HYBRID_RRF" and "0.831" become column NAMES. A model reading that gets
    # every remaining value bound to a nonsense label, in flawless markdown, with
    # a real page citation.
    #
    # The repair is the second extractor that already exists. page.get_text()
    # renders the region verbatim -- it loses the tidy columns and keeps every
    # number next to its actual label. Built as a detector this morning; it is
    # also a repair path, because detection tells you which representation is
    # broken and the disagreeing extractor supplies another one.
    #
    # CLIPPED TO THE TABLE BBOX rather than taking the whole page: a two-column
    # page is 600-900 tokens against a ~100-token table chunk, and at a
    # 3000-token context budget one repaired table would eat a third of it.
    for i in range(doc.page_count):
        page = doc[i]
        for t in per_page[i]:
            if t.header_ok:
                continue
            x0, y0, x1, y1 = t.bbox
            rect = pymupdf.Rect(x0 - 4, y0 - 4, x1 + 4, y1 + 4)
            t.clip = (page.get_text(clip=rect) or "").strip()
    doc.close()


    # -- text pass ------------------------------------------------------------
    pages = pymupdf4llm.to_markdown(str(path), page_chunks=True, show_progress=False)
    blocks: list[Block] = []
    parsed_parts: list[str] = []
    ordinal = 0
    n_repaired = 0
    n_repair_unverified = 0
    for pg in pages:
        # pymupdf4llm names this page_number. Reading "page" returned None ->
        # 0 for every chunk, so the suspect-page set never matched and the flag
        # silently never fired: 1 suspect detected, 0 blocks flagged. The
        # detector was right and the plumbing dropped it on the floor -- which is
        # why the test asserted on the FLAGGED BLOCK COUNT and not just on the
        # suspect count.
        meta = pg.get("metadata", {}) or {}
        page_no = int(meta.get("page_number") or meta.get("page") or 0)
        md = pg.get("text", "") or ""
        parsed_parts.append(md)
        # One matched pass: the block's flags AND its repair both come from the
        # single candidate this block was matched to. Unmatched table blocks get
        # no flags at all rather than inheriting their page's -- refusing to
        # attribute a defect we cannot locate beats over-flagging, now that
        # something downstream acts on the flags.
        cands = list(per_page[page_no - 1]) if 1 <= page_no <= len(per_page) else []
        for blk in _split_markdown_blocks(md):
            text = blk["text"]
            source = "markdown"
            hdr_missing = False
            continuation = False
            if blk["kind"] is ChunkKind.TABLE and cands:
                m = match_candidate(blk["text"], cands)
                if m is not None:
                    cands.remove(m)
                    hdr_missing = not m.header_ok
                    continuation = m.is_continuation_suspect
                    if m.clip:
                        text = m.clip
                        source = "page_text_clip"
                        n_repaired += 1
                else:
                    n_repair_unverified += 1
            blocks.append(
                Block(
                    kind=blk["kind"],
                    text=text,
                    text_source=source,
                    page=page_no,
                    heading_path=blk["heading_path"],
                    ordinal=ordinal,
                    # Page-level granularity, and that is coarser than ideal:
                    # every block on a suspect page is flagged, not only the
                    # table. Honest for a detector whose job is to raise a hand,
                    # and it errs toward over-flagging, which is the safe
                    # direction when nothing is being modified.
                    table_continuation_suspect=continuation,
                    table_header_missing=hdr_missing,
                )
            )
            ordinal += 1

    # Hash of the PARSER OUTPUT. A parser that drifts between runs leaves
    # content_hash and the pipeline fingerprint identical while the indexed text
    # differs, so provenance would certify two different corpora as one.
    src.parsed_hash = sha256_text("\n".join(parsed_parts))

    diagnostics = {
        "n_pages": src.n_pages,
        "n_tables_detected": len(all_tables),
        "n_tables_headerless": len({(t.page, t.bbox) for t in all_tables if not t.header_ok}),
        "n_continuation_suspects": len(suspects),
        "suspects": suspects,
        "detector": "find_tables(strategy='lines')",
        "detector_blind_spot": "unruled/whitespace-aligned tables are not detected at all",
        "n_tables_repaired_page_text": n_repaired,
        "n_repairs_unverified": n_repair_unverified,
        "repair_note": (
            "header-missing tables are delivered as bbox-clipped page text, not as "
            "markdown -- the markdown promotes a data row to a column header"
        ),
    }
    return src, blocks, diagnostics


_H = re.compile(r"^(#{1,6})\s+(.*)$")
# pymupdf4llm emits headings as "## **3 Results**" -- bold markers included. Left
# in, those markers ride into the breadcrumb and therefore into embed_text, so we
# would be embedding markdown syntax as though it were content. Cheap to strip,
# and it matters because the breadcrumb is a large share of a small chunk's
# tokens (measured: 9% overhead on a paper, 39% on manual.docx).
_MD_EMPH = re.compile(r"(\*\*|__|\*|_|`)")


# Minimum alphanumeric characters for a heading to count as one. graphrag's PDF
# yields a heading of '{{' -- a LaTeX artifact. Two characters, no letters. Left
# in, it becomes a breadcrumb that is prepended to every child of its section and
# embedded as content, and it triggered a spurious prefix-leak assertion because
# '{{' occurs in ordinary body text by coincidence.
_MIN_HEADING_ALNUM = 3


def clean_heading(text: str) -> str:
    out = " ".join(_MD_EMPH.sub("", text).split()).strip(" :#-|{}")
    if sum(c.isalnum() for c in out) < _MIN_HEADING_ALNUM:
        return ""
    return out


def _split_markdown_blocks(md: str) -> list[dict[str, Any]]:
    """Split markdown into blocks while carrying the heading breadcrumb down.

    heading_path is not decoration. It is the cheapest possible version of
    Contextual Retrieval (M5): 'Recall rose 8 points' is unretrievable alone, but
    ('3 Results', '3.2 Retrieval quality') tells you what it is about at zero API
    cost. Session 4 adds the LLM-written prefix on top; this is the free part,
    and measuring the free part first is how you find out whether the paid part
    was worth it.
    """
    out: list[dict[str, Any]] = []
    heading: list[str] = []
    buf: list[str] = []

    def flush(kind: ChunkKind = ChunkKind.TEXT) -> None:
        text = "\n".join(buf).strip()
        buf.clear()
        if text:
            out.append({"kind": kind, "text": text, "heading_path": tuple(heading)})

    for line in md.split("\n"):
        m = _H.match(line.strip())
        if m:
            flush()
            level = len(m.group(1))
            cleaned = clean_heading(m.group(2))
            # A rejected heading truncates the path rather than inserting an
            # empty level: an empty string in heading_path would render as a
            # stray " > " in every breadcrumb below it.
            heading = heading[: level - 1] + ([cleaned] if cleaned else [])
            continue
        if line.strip().startswith("|") and buf and not buf[-1].strip().startswith("|"):
            flush()  # a markdown table starts: keep it as its own block
        if buf and buf[-1].strip().startswith("|") and not line.strip().startswith("|"):
            flush(ChunkKind.TABLE)
        buf.append(line)
    flush(ChunkKind.TABLE if buf and buf[0].strip().startswith("|") else ChunkKind.TEXT)
    return out


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def load_docx(path: Path) -> tuple[Source, list[Block], dict[str, Any]]:
    """DOCX is the control case: heading structure is DECLARED, not inferred.

    A PDF encodes glyph positions, so headings must be guessed from font size and
    weight. A DOCX says 'Heading 2'. That makes it the baseline against which PDF
    heading detection can be judged -- if header-aware chunking helps on DOCX and
    not on PDF, the problem is heading inference, not chunking.
    """
    from docx import Document as Docx
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph

    src = Source.from_file(path, DocType.DOCX)
    d = Docx(str(path))

    blocks: list[Block] = []
    heading: list[str] = []
    ordinal = 0
    parts: list[str] = []

    # Walk body children in document order so tables land between the right
    # paragraphs. Iterating d.paragraphs then d.tables would lose that ordering
    # and put every table at the end, detaching each one from its section.
    body = d.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            para = Paragraph(child, d)
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "") if para.style else ""
            m = re.match(r"Heading (\d+)", style)
            if m or style == "Title":
                level = 0 if style == "Title" else int(m.group(1))  # type: ignore[union-attr]
                heading = heading[:level] + [text] if level else [text]
                parts.append(("#" * max(level, 1)) + " " + text)
                continue
            blocks.append(
                Block(kind=ChunkKind.TEXT, text=text, heading_path=tuple(heading), ordinal=ordinal)
            )
            parts.append(text)
            ordinal += 1
        elif tag == "tbl":
            tbl = DocxTable(child, d)
            rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
            ok, why = header_plausible(rows)
            md = _rows_to_markdown(rows)
            blocks.append(
                Block(
                    kind=ChunkKind.TABLE,
                    text=md,
                    heading_path=tuple(heading),
                    ordinal=ordinal,
                    table_header_missing=not ok,
                )
            )
            parts.append(md)
            ordinal += 1

    src.parsed_hash = sha256_text("\n".join(parts))
    n_tbl = sum(1 for b in blocks if b.kind is ChunkKind.TABLE)
    return src, blocks, {
        "n_pages": 0,  # DOCX has no pages until it is rendered
        "n_tables_detected": n_tbl,
        "n_tables_headerless": sum(1 for b in blocks if b.table_header_missing),
        "n_continuation_suspects": 0,  # no pagination, so no page-break split
        "suspects": [],
        "detector": "docx native table objects (structure is declared)",
        "detector_blind_spot": "none for tables; DOCX declares them",
    }


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [list(r) + [""] * (width - len(r)) for r in rows]
    head = "| " + " | ".join(norm[0]) + " |"
    rule = "| " + " | ".join("---" for _ in range(width)) + " |"
    body = ["| " + " | ".join(r) + " |" for r in norm[1:]]
    return "\n".join([head, rule, *body])


# --------------------------------------------------------------------------
# Images
# --------------------------------------------------------------------------

_MIME = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"}


def load_image(path: Path, *, caption: bool = False) -> tuple[Source, list[Block], dict[str, Any]]:
    """An image with no caption is UNRETRIEVABLE, not merely worse.

    Nothing in an index can match a query against pixels, so the caption is the
    only thing that makes the image findable. That is why a skipped caption is a
    retrieval failure and gets a degradation notice rather than a shrug.

    `caption=False` by default so that walking the corpus does not silently spend
    a free-tier quota that Session 4 needs. The caller opts in.
    """
    src = Source.from_file(path, DocType.IMAGE)
    mime = _MIME.get(path.suffix.lower(), "image/png")

    if not caption:
        src.parsed_hash = sha256_text(f"uncaptioned:{src.content_hash}")
        return src, [], {
            "n_pages": 1, "n_tables_detected": 0, "n_tables_headerless": 0,
            "n_continuation_suspects": 0, "suspects": [],
            "detector": "none (captioning not requested)",
            "detector_blind_spot": "image is present but absent from the retrieval index",
        }

    from .. import gemini

    data = path.read_bytes()
    ctx = path.stem.replace("_", " ")
    limits.seen("caption_image", n=1)
    try:
        text, _usage = gemini.caption_image(data, mime, context=ctx)
    except (limits.QuotaExhausted, gemini.EmptyResponse):
        # THE NOTICE NOW ACTUALLY FIRES. limits.caption_skipped() existed from the
        # start and nothing ever called it -- so a captioning failure produced an
        # image with no caption, silently, and an image with no caption is
        # UNRETRIEVABLE rather than merely degraded. Found by the reachability
        # audit, not by anything failing.
        limits.report(limits.caption_skipped(1, 1))
        src.parsed_hash = sha256_text(f"caption_failed:{src.content_hash}")
        return src, [], {
            "n_pages": 1, "n_tables_detected": 0, "n_tables_headerless": 0,
            "n_continuation_suspects": 0, "suspects": [],
            "detector": "gemini vision captioning FAILED",
            "detector_blind_spot": "image is present but absent from the retrieval index",
        }
    src.parsed_hash = sha256_text(text)
    return src, [
        Block(
            kind=ChunkKind.IMAGE_CAPTION,
            text=text,
            text_source="gemini_caption",  # nobody wrote this; a model did
            page=1,
            asset_path=str(path),
            ordinal=0,
        )
    ], {
        "n_pages": 1, "n_tables_detected": 0, "n_tables_headerless": 0,
        "n_continuation_suspects": 0, "suspects": [],
        "detector": "gemini vision captioning",
        "detector_blind_spot": "caption is model-written; nothing in it is verbatim",
    }


# --------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------

_LOADERS = {
    ".pdf": DocType.PDF,
    ".docx": DocType.DOCX,
    ".png": DocType.IMAGE,
    ".jpg": DocType.IMAGE,
    ".jpeg": DocType.IMAGE,
    ".webp": DocType.IMAGE,
}


def load(
    path: Path, *, caption_images: bool = False, use_cache: bool = True
) -> tuple[Source, list[Block], dict[str, Any]]:
    kind = _LOADERS.get(path.suffix.lower())
    if kind is None:
        raise ValueError(f"no loader for {path.suffix} ({path.name})")

    if use_cache and kind is not DocType.IMAGE:
        probe = Source.from_file(path, kind)
        hit = _cache_load(probe)
        if hit is not None:
            return hit

    if kind is DocType.PDF:
        src, blocks, diag = load_pdf(path)
    elif kind is DocType.DOCX:
        src, blocks, diag = load_docx(path)
    else:
        # Images are NOT parse-cached here. The expensive part is a Gemini call,
        # and it belongs in the embedding/LLM cache with its own provenance --
        # caching a caption under "parse" would hide a model-generated artifact
        # inside something named after deterministic parsing.
        return load_image(path, caption=caption_images)

    diag["from_cache"] = False
    if use_cache:
        _cache_store(src, blocks, diag)
    return src, blocks, diag


def corpus_files(root: Path | None = None) -> list[Path]:
    root = root or config.DATA_RAW
    return sorted(p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in _LOADERS)
