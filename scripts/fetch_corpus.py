"""
Build the corpus. Idempotent: run it twice, download nothing the second time.

WHY THESE DOCUMENTS. Three properties, each earning its place:

1. THEY BREAK NAIVE PARSING. arXiv papers are two-column with tables that span
   page breaks, figures with captions, footnotes, and running headers. Naive
   extract_text() interleaves the columns -- which is the point. A corpus that
   parses cleanly teaches nothing about parsing.

2. THEY CARRY EXACT IDENTIFIERS. Citation keys, equation numbers, arXiv ids,
   hyperparameter names (efConstruction, k1, RRF k=60). Session 5 needs queries
   containing exact tokens to show the dense-vs-BM25 split, and dense retrieval
   fails on precisely this class of token.

3. YOU HAVE READ THEM. They are your guide's own reading list. When the chatbot
   answers badly you will know WHY, instead of only knowing that it did. Judging
   a RAG system over a corpus you do not know is how people end up trusting
   demos -- you cannot distinguish a retrieval failure from a corpus that never
   contained the answer.

Plus two synthetic additions, and the reason each is synthetic:

   manual.docx -- a DOCX with real heading structure and INVENTED error codes.
     Synthetic on purpose: the identifier-query experiment needs GROUND TRUTH.
     If I mine identifiers out of the papers I have to verify each one by hand;
     if I author them, I know exactly which chunk contains ERR_5521 and can
     compute recall without labelling anything.

   benchmark-report.pdf -- a table whose header is on page 1 and whose data rows
     are on page 2, with no repeated header. Synthetic because a survey of the
     eight papers found exactly one table touching a bottom margin, and none
     with a known-correct answer to check a stitch against. Here the correct
     answer to "HYBRID_RRF's Recall@10" is 0.831 and it sits on page 2
     unlabelled -- so a pipeline that answers 0.712 has reproduced the
     quietly-incomplete failure, checkably, with no labelling work.

   assets/*.png -- pages rendered as images, for the vision path. Real figures
     from real papers, so the captioning model has something with axes, labels
     and numbers to transcribe rather than clip art.
"""

from __future__ import annotations

import sys
import time
from pathlib import Path

import httpx

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from ragkit import config  # noqa: E402

# From the guide's Reading Order section.
PAPERS: list[tuple[str, str]] = [
    ("2307.03172", "lost-in-the-middle"),        # Liu et al. 2023 -- tables, long ctx
    ("2401.05856", "seven-failure-points"),      # Barnett et al. 2024 -- the taxonomy
    ("2212.10496", "hyde"),                      # Gao et al. 2022
    ("2401.18059", "raptor"),                    # Sarthi et al. 2024 -- dense tables
    ("2404.16130", "graphrag-local-to-global"),  # Edge et al. 2024
    ("1603.09320", "hnsw"),                      # Malkov & Yashunin -- heavy math/figures
    ("2309.15217", "ragas"),                     # Es et al. 2023
    ("2310.11511", "self-rag"),                  # Asai et al. 2023
]

# Pages rendered to PNG for the vision path: (paper slug, 1-based page numbers).
# Chosen as pages that carry figures or tables rather than prose.
RENDER_PAGES: list[tuple[str, tuple[int, ...]]] = [
    ("lost-in-the-middle", (1, 5)),
    ("hnsw", (2,)),
    ("raptor", (2,)),
]


def fetch_papers(dest: Path) -> tuple[int, int]:
    dest.mkdir(parents=True, exist_ok=True)
    got = skipped = 0
    with httpx.Client(
        follow_redirects=True,
        timeout=60.0,
        headers={"User-Agent": "learning-rag/0.1 (educational RAG build)"},
    ) as cl:
        for arxiv_id, slug in PAPERS:
            out = dest / f"{slug}.pdf"
            # Idempotency, with a size sanity check: a 2KB "PDF" is an arXiv
            # error page, and silently indexing it would look like a bad parser.
            if out.exists() and out.stat().st_size > 50_000:
                print(f"  skip  {out.name} ({out.stat().st_size // 1024} KB)")
                skipped += 1
                continue
            url = f"https://arxiv.org/pdf/{arxiv_id}"
            try:
                r = cl.get(url)
                r.raise_for_status()
                if not r.content.startswith(b"%PDF"):
                    print(f"  FAIL  {slug}: response is not a PDF ({len(r.content)} bytes)")
                    continue
                out.write_bytes(r.content)
                print(f"  got   {out.name} ({len(r.content) // 1024} KB)")
                got += 1
            except Exception as exc:  # noqa: BLE001
                print(f"  FAIL  {slug}: {type(exc).__name__}: {exc}")
            time.sleep(1.0)  # arXiv asks for this; do not hammer a free service
    return got, skipped


def write_manual_docx(dest: Path) -> Path:
    """A DOCX with genuine heading levels and invented identifiers.

    Heading levels matter: header-aware chunking (Session 3) splits on document
    structure, and a DOCX is the one format where that structure is unambiguous
    -- python-docx gives you 'Heading 1' / 'Heading 2' directly, with no layout
    inference. It is the control case against which PDF heading detection can be
    judged.
    """
    from docx import Document as Docx

    d = Docx()
    d.add_heading("Meridian Retrieval Appliance - Service Manual", level=0)
    d.add_paragraph(
        "Applies to appliance models MRA-4200, MRA-4400 and MRA-8800, "
        "firmware 7.3 and later. Document revision D-2291."
    )

    d.add_heading("1. Diagnostic Error Codes", level=1)
    d.add_paragraph(
        "The appliance reports faults as an alphanumeric code on the front panel "
        "display. Codes persist in the event log for 90 days."
    )

    d.add_heading("1.1 Index Subsystem", level=2)
    for code, meaning, action in [
        ("ERR_5521", "Vector index checksum mismatch after an unclean shutdown.",
         "Rebuild the index with 'mra-admin index --rebuild'. Expect 40 minutes per million vectors."),
        ("ERR_5522", "Index shard unreachable; the shard process exited without flushing.",
         "Restart the shard. If ERR_5522 recurs within one hour, replace part number PN-8891-A."),
        ("ERR_5530", "Embedding dimension mismatch between the index and the configured model.",
         "The index was built at a different dimensionality. Re-embed the corpus; do not edit the config."),
    ]:
        d.add_paragraph(f"{code} - {meaning}", style="List Bullet")
        d.add_paragraph(f"Corrective action: {action}")

    d.add_heading("1.2 Query Subsystem", level=2)
    for code, meaning, action in [
        ("ERR_6104", "Query latency exceeded the configured p99 budget of 800 ms.",
         "Lower efSearch from 256 to 128 and re-measure. Recall will fall by roughly two points."),
        ("ERR_6110", "Reranker model failed to load.",
         "Verify the model cache at /var/lib/mra/models. Free at least 4 GB."),
    ]:
        d.add_paragraph(f"{code} - {meaning}", style="List Bullet")
        d.add_paragraph(f"Corrective action: {action}")

    d.add_heading("2. Field-Replaceable Parts", level=1)
    d.add_paragraph(
        "Order parts by the exact part number. Substitutions void the warranty."
    )
    table = d.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Part Number", "Description", "Models", "Lead Time"]):
        hdr[i].text = h
    for row in [
        ("PN-8891-A", "Index shard controller board", "MRA-4400, MRA-8800", "3 days"),
        ("PN-8891-B", "Index shard controller board, revision B", "MRA-8800 only", "10 days"),
        ("PN-7712-C", "Front panel display assembly", "all models", "1 day"),
        ("PN-9004-F", "Cooling fan module, 80 mm", "MRA-4200, MRA-4400", "same day"),
    ]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    d.add_heading("3. Maintenance Schedule", level=1)
    d.add_heading("3.1 Quarterly", level=2)
    d.add_paragraph(
        "Verify index recall against the stored ground-truth query set. A drop "
        "of more than three points indicates corpus drift, not hardware fault."
    )
    d.add_heading("3.2 Annually", level=2)
    d.add_paragraph(
        "Replace PN-9004-F regardless of apparent condition. Bearing failure is "
        "not detectable before it occurs."
    )

    out = dest / "manual.docx"
    d.save(str(out))
    print(f"  wrote {out.name}")
    return out


def write_spanning_table_pdf(dest: Path) -> Path:
    """A PDF whose table header is on page 1 and whose data rows are on page 2.

    SYNTHETIC for the same reason manual.docx is: GROUND TRUTH. A survey of the
    eight arXiv papers found exactly ONE table touching a bottom margin, and none
    with a known-correct answer to check a stitch against.

    The authored defect:
      - header row is  Method | Recall@10 | Precision@5 | Latency_ms
      - rows 1-3 sit at the bottom of page 1
      - rows 4-6 sit at the top of page 2, with NO repeated header
      - correct answer to "HYBRID_RRF Recall@10" is 0.831, on page 2, unlabelled

    That is the test case: a pipeline answering 0.712 (the best value visible on
    page 1) has reproduced the quietly-incomplete failure, checkably, with no
    labelling work.

    RULING LINES ARE DRAWN DELIBERATELY, and version one of this file did not
    draw them. That was a bad fixture, and finding out why was worth more than
    the fixture: PyMuPDF find_tables() defaults to strategy="lines", so a
    whitespace-aligned table has NO ruling lines and is not detected as a table
    at all -- it silently becomes prose, and the column alignment that carried
    the row-column binding is destroyed by chunking. Meanwhile strategy="text"
    detects tables everywhere, returning page 1's title and prose as a 23x3
    table with words split mid-token.

    So detection fails in two opposite directions before any table STRATEGY
    (markdown / summary / real-table) gets a chance to run, and both of those
    strategies presuppose detection succeeded. Table 5 below is kept
    whitespace-only on purpose, as the fixture for that stage-zero case.

    Table 4 has a legitimately non-numeric header, so a header check can be
    tested for FALSE POSITIVES too. A validator that rejects every table is as
    useless as one that accepts every table.
    """
    import pymupdf

    mono = "cour"
    COLS = [60, 210, 320, 430, 530]   # 4 columns -> 5 boundaries
    RH = 18                            # row height

    def grid(page, top: int, rows: list[list[str]]) -> None:
        """Draw text plus a ruled grid, so strategy='lines' can see a table."""
        for r, cells in enumerate(rows):
            y = top + r * RH
            for c, val in enumerate(cells):
                # NOTE, and my first explanation of this was wrong. Table
                # cells extract as "HYBRID RRF" plus a stray underscore on
                # its own line, while page.get_text() on the SAME page
                # returns HYBRID_RRF verbatim. I blamed the Courier font and
                # switched to helv; that changed nothing. The mangling lives
                # in PyMuPDF find_tables().extract(), not in the font or the
                # text layer.
                #
                # Consequence for Session 5, which is the reason this note
                # exists: indexing table cells straight from extract() would
                # destroy exactly the tokens BM25 is there to match. The
                # loader has to reconcile cell text against get_text(), or
                # treat cell text as lower-fidelity than page text.
                page.insert_text((COLS[c] + 4, y + 13), val, fontsize=9, fontname="helv")
        bottom = top + len(rows) * RH
        for r in range(len(rows) + 1):                      # horizontals
            y = top + r * RH
            page.draw_line(pymupdf.Point(COLS[0], y), pymupdf.Point(COLS[-1], y), width=0.6)
        for x in COLS:                                      # verticals
            page.draw_line(pymupdf.Point(x, top), pymupdf.Point(x, bottom), width=0.6)

    doc = pymupdf.open()

    p1 = doc.new_page()
    y = 60
    for line, size in [
        ("Meridian Retrieval Benchmark, Internal Report R-4471", 13),
        ("", 10),
        ("3. Results", 11),
        ("", 10),
        ("Table 3 reports retrieval quality across six configurations. The", 10),
        ("evaluation set contains 400 queries, 90 of which contain an exact", 10),
        ("identifier such as a part number or an error code.", 10),
        ("", 10),
        ("Table 3: Retrieval quality by configuration.", 10),
    ]:
        p1.insert_text((60, y), line, fontsize=size, fontname=mono)
        y += size + 6

    # header + rows 1-3, pushed to the bottom of page 1
    grid(p1, 700, [
        ["Method", "Recall@10", "Precision@5", "Latency_ms"],
        ["DENSE_ONLY", "0.684", "0.512", "46"],
        ["BM25_ONLY", "0.601", "0.549", "11"],
        ["HYBRID_ALPHA", "0.712", "0.573", "52"],
    ])

    # continuation: identical columns, identical x-positions, NO header
    p2 = doc.new_page()
    grid(p2, 60, [
        ["HYBRID_RRF", "0.831", "0.604", "58"],
        ["HYBRID_RRF_RR", "0.889", "0.702", "318"],
        ["COLBERT_LATE", "0.874", "0.681", "141"],
    ])

    y = 160
    for line in [
        "HYBRID_RRF fuses the dense and sparse legs with reciprocal rank",
        "fusion at k=60. HYBRID_RRF_RR adds a cross-encoder reranker over",
        "the top 50 candidates, which is where the latency comes from.",
        "",
        "Table 4: Configuration parameters.",
    ]:
        p2.insert_text((60, y), line, fontsize=10, fontname=mono)
        y += 16

    grid(p2, y + 6, [
        ["Parameter", "Value", "Applies at", "Notes"],
        ["efSearch", "128", "query time", "runtime recall dial"],
        ["efConstruction", "200", "build time", "no query effect"],
        ["M", "32", "build time", "edges per node"],
        ["rrf_k", "60", "fusion", "rank smoothing"],
    ])

    # Table 5: whitespace-aligned, NO ruling lines -- the stage-zero fixture.
    y2 = y + 6 + 5 * RH + 40
    for line in [
        "Table 5: Legacy results, retained for reference (no ruling lines).",
        "",
        "Method          Recall@10   Notes",
        "LEGACY_TFIDF    0.412       superseded by BM25_ONLY",
        "LEGACY_LSA      0.388       superseded by DENSE_ONLY",
    ]:
        p2.insert_text((60, y2), line, fontsize=9, fontname=mono)
        y2 += 14

    out = dest / "benchmark-report.pdf"
    doc.save(str(out))
    doc.close()
    print(f"  wrote {out.name} (T3 header p1 / rows p2; T4 ruled; T5 unruled)")
    return out


def render_page_images(dest: Path) -> int:
    """Render selected PDF pages to PNG for the vision/captioning path."""
    import pymupdf

    assets = dest / "assets"
    assets.mkdir(parents=True, exist_ok=True)
    n = 0
    for slug, pages in RENDER_PAGES:
        pdf = dest / f"{slug}.pdf"
        if not pdf.exists():
            continue
        doc = pymupdf.open(pdf)
        for pno in pages:
            if pno > doc.page_count:
                continue
            out = assets / f"{slug}_p{pno}.png"
            if out.exists():
                continue
            # 150 dpi: legible axis labels and table digits for the captioner
            # without producing multi-megabyte payloads on a free key.
            pix = doc[pno - 1].get_pixmap(dpi=150)
            pix.save(str(out))
            print(f"  wrote assets/{out.name} ({out.stat().st_size // 1024} KB)")
            n += 1
        doc.close()
    return n


def main() -> None:
    config.ensure_dirs()
    dest = config.DATA_RAW
    print(f"corpus -> {dest}\n")
    print("PDFs (arXiv, from the guide's reading list):")
    got, skipped = fetch_papers(dest)
    print("\nDOCX (synthetic, for identifier ground truth):")
    write_manual_docx(dest)
    print("\nPDF (synthetic, for the page-spanning table case):")
    write_spanning_table_pdf(dest)
    print("\nPage images (for the vision path):")
    n_img = render_page_images(dest)

    pdfs = sorted(dest.glob("*.pdf"))
    imgs = sorted((dest / "assets").glob("*.png"))
    total_mb = sum(p.stat().st_size for p in dest.rglob("*") if p.is_file()) / 1e6
    print(
        f"\nCORPUS: {len(pdfs)} PDFs, 1 DOCX, {len(imgs)} images, {total_mb:.1f} MB"
        f"  (downloaded {got}, already present {skipped}, rendered {n_img})"
    )
    if len(pdfs) < 4:
        print(
            "\nWARNING: fewer than 4 PDFs. The parser comparison and the "
            "dense-vs-BM25 split both need a real corpus to show anything. "
            "Re-run, or drop your own PDFs into data/raw/."
        )


if __name__ == "__main__":
    main()
